import os

import win32api
import win32con
import win32process
from docx import Document
from win32com import client
from win32com.client import DispatchEx
import pandas as pd

def close_excel_by_force(excel):  # 关闭进程

    # Get the window's process id's
    hwnd = excel.Hwnd
    t, p = win32process.GetWindowThreadProcessId(hwnd)
    # Ask window nicely to close
    try:
        handle = win32api.OpenProcess(win32con.PROCESS_TERMINATE, 0, p)
        if handle:
            win32api.TerminateProcess(handle, 0)
            win32api.CloseHandle(handle)
    except:
        pass

def save_doc_to_docx(rawpath, namestring):#doc 化为docx
    #模拟打开word
    word = client.Dispatch("Word.Application")
    word.Visible = 0 #后台运行
    word.DisplayAlerts = 0 #不显示不警告
    print(rawpath+'/'+namestring)
    doc = word.Documents.Open(rawpath+'/'+namestring)
    # 将文件名与后缀分
    rename = os.path.splitext(namestring)
    #将文件另存为'docx', 12代表docx
    # print(rawpath+'/'+rename[0]+'.docx')
    doc.SaveAs(rawpath+'/'+rename[0]+'.docx', 12)
    doc.Close()
    word.Quit()
    close_excel_by_force(DispatchEx("Excel.Application"))
    return rename[0]+'.docx'

def change_file_name(rawpath, namestring):#doc 化为docx
    #模拟打开word
    word = client.Dispatch("Word.Application")
    word.Visible = 0 #后台运行
    word.DisplayAlerts = 0 #不显示不警告
    print(rawpath+'/'+namestring)
    doc = word.Documents.Open(rawpath+'/' + namestring)
    # 将文件名与后缀分
    document = Document(rawpath + '/' + namestring)
    table = document.tables  # 读取word中的表格
    table0 = table[0]
    # rename = os.path.splitext(namestring)
    #将文件另存为'docx', 12代表docx
    # print(rawpath+'/'+rename[0]+'.docx')
    doc.SaveAs(rawpath+'/' + table0.cell(0, 1).text + '.docx', 12)
    doc.Close()
    word.Quit()
    close_excel_by_force(DispatchEx("Excel.Application"))

fileList = []
path = 'C:/Users/86410/Desktop/试验数据/试验报告'
# for i in range(1, 8):
#     print(path+str(i))
#     for file_name in os.listdir(path+str(i)):
#         if file_name[-3:] == 'doc':
#             # print(file_name)
#             save_doc_to_docx(path+str(i), file_name)
#             os.remove(os.path.join(path+str(i), file_name))
# print('保存完毕')
#############################################
# for i in range(1, 8):
#     print(path+str(i))
#     for file_name in os.listdir(path+str(i)):
#         change_file_name(path+str(i), file_name)
#         os.remove(os.path.join(path+str(i), file_name))
# print('保存完毕')
###############################################
# for file_name in os.listdir(path):
#     document = Document(path + '/' + file_name)
#     table = document.tables  # 读取word中的表格
#     document.
#     print(file_name+":" + table.info())
des_path = 'C:/Users/86410/Desktop/试验数据/最后数据'


# for i in range(1, 8):
#     print(path+str(i))
#     for file_name in os.listdir(path+str(i)):
#         document = Document(path+str(i) + '/' + file_name)
#         tables = document.tables
#         first_table = tables[0]
#         place = ''
#         time = ''
#         cell_row = 0
#         for row in first_table.rows:
#             cell_column = 0
#             for cell in row.cells:
#                 if cell.text == '工作地点':
#                     place = row.cells[cell_column+1].text.strip().split(',')[0]
#                     print(place)
#                 elif cell.text == '工作开始时间':
#                     time = row.cells[cell_column+1].text.split('-')[0]
#                     print(time)
#                 cell_column += 1
#             cell_row += 1
#
#         for table in tables[1:]:
#             item_list = []
#             tan_list = []
#             change_dr = []
#             if table.cell(0, 1).text == 'tanδ(%)' and len(table.cell(1, 1).text)>0:
#                 # doc = Document()
#                 # doc.add_table(table[count])
#                 # doc.save(des_path + '/' + file_name)
#                 for row in table.rows[1:-1]:
#                     item_list.append(row.cells[0].text)
#                     tan_list.append(row.cells[1].text)
#                     change_dr.append(row.cells[4].text)
#
#                 data = pd.DataFrame({'变电站': place, '所属相别': item_list, 'tg(%)': tan_list, '电容变化': change_dr, '试验日期': time})
#                 # data = data.drop(0, axis=1)
#                 data.to_csv(des_path + '/' + file_name.split('.')[0]+'.csv', encoding='gbk', index=False)
#                 break
# print('保存完毕')
res_file_list = []
res_file_list = os.listdir(des_path)
res_dic = {'变电站': '', '所属相别': '', 'tg(%)': '', '电容变化': '', '试验日期': '', 'tg(%)2': '', '电容变化2': '', '试验日期2': '', 'tg(%)3': '', '电容变化3': '', '试验日期3': ''}
res_list = []
finish_list = []
for item in res_file_list:
    # print(des_path+'/'+item)
    data = pd.read_csv(des_path + '/' + item, encoding='gbk') #逐个读取csv文件
    # res_dic.get('试验日期').append(i)
    if item.split('.')[0].split('(')[0].strip() not in finish_list:
        for i in range(len(data)):
            res_dic['变电站'] = data.loc[i]['变电站']
            res_dic['所属相别'] = data.loc[i]['所属相别']
            res_dic['tg(%)'] = data.loc[i]['tg(%)']
            res_dic['电容变化'] = data.loc[i]['电容变化']
            res_dic['试验日期'] = data.loc[i]['试验日期']
            res_list.append(res_dic.copy())
    else:
        for index, i in enumerate(res_list):
            try:
                if index == 608:
                    print('a')
                if i.get('变电站') == data.loc[0]['变电站'] and i.get('所属相别') == '高压套管-A':
                    if res_list[index]['tg(%)2'] == '':
                        for j in range(len(data)):
                            res_list[index+j]['tg(%)2'] = data.loc[j]['tg(%)']
                            res_list[index+j]['电容变化2'] = data.loc[j]['电容变化']
                            res_list[index+j]['试验日期2'] = data.loc[j]['试验日期']

                    elif res_list[index]['tg(%)3'] == '':
                        for j in range(len(data)):
                            res_list[index + j]['tg(%)3'] = data.loc[j]['tg(%)']
                            res_list[index + j]['电容变化3'] = data.loc[j]['电容变化']
                            res_list[index + j]['试验日期3'] = data.loc[j]['试验日期']
                        # break
            except Exception as e:
                print(index)
    finish_list.append(item.split('.')[0].split('(')[0].strip())
    # res = pd.DataFrame({'变电站': place, '所属相别': item_list, 'tg(%)': tan_list, '电容变化': change_dr, '试验日期': time})
pd.DataFrame(res_list).to_csv(des_path+'/整合数据.csv', index=False, encoding='gbk')